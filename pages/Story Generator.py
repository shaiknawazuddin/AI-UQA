import streamlit as st
import pandas as pd
import openai
from docx import Document
import zipfile
import os
import time
from io import BytesIO


st.markdown("""
<style>
.st-emotion-cache-6q9sum.ef3psqc3
{
            visibility: hidden;
}
</style>
""", unsafe_allow_html = True)

st.markdown("""
<style>
.st-emotion-cache-cio0dv.ea3mdgi1
{
            visibility: hidden;
}
</style>
""", unsafe_allow_html = True)

st.markdown("""
<style>
.st-emotion-cache-ch5dnh.ef3psqc4
{
            visibility: hidden;
}
</style>
""", unsafe_allow_html = True)

# Get the OpenAI API key from the user via a text input
openai_api_key = st.text_input("Enter your OpenAI API key")

# Check if the API key is provided
if not openai_api_key:
    st.warning("Please enter your OpenAI API key.")
    st.stop()

# Set the OpenAI API key
openai.api_key = openai_api_key

st.title("Story Generator :thought_balloon:")



st.sidebar.title('Input method')

def save_word_docx(filename, content):
    doc = Document()
    title_list = str(filename).split('_')
    title = ' '.join(title_list)

    doc.add_heading(title, 0)

    # Convert content to utf-8 to handle special characters
    content_utf = content.encode('latin-1', 'replace').decode('utf-8')

    doc.add_paragraph(content_utf)

    docx_file = f"{filename}.docx"
    doc.save(docx_file)

    return docx_file

# Backend Code 1 -- End

def quran_portion():
    str_input = st.sidebar.text_input("Enter a Quran Portion (e.g., a verse or a theme)")
    gpt_version = st.sidebar.radio("Select GPT Version", ["GPT 3.5", "GPT 4"])
    df = pd.DataFrame({'Column 1' : [str_input]})
    custom_prompt = st.sidebar.text_area("Enter a Custom Prompt (optional)")
    return gpt_version, df, custom_prompt

def upload_and_process_file():
    uploaded_file = st.sidebar.file_uploader("Upload an Excel/CSV file", type=['csv', 'xlsx'])
    custom_prompt = st.sidebar.text_area("Enter a Custom Prompt (optional)")
    gpt_version = st.sidebar.radio("Select GPT Version", ["GPT 3.5", "GPT 4"])

    if uploaded_file is not None:
        if uploaded_file.type == 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet':
            df = pd.read_excel(uploaded_file)
        else:
            df = pd.read_csv(uploaded_file)
        return gpt_version, df, custom_prompt
    return None, None, None

# Backend Code 2 -- Start

def output_code(gpt_version, df):
    code_start = time.time()
    if df is not None:
        if gpt_version == "GPT 3.5":
            model = "gpt-3.5-turbo"
        else:
            model = "gpt-4"
        story_list = []
        df = df.rename(columns = {df.columns[0] : 'Surah Number'})
        docx_files = []  # Keep track of generated Word files

        for index, row in df.iterrows():
            theme = row['Surah Number']  # Assuming 'Column 1' contains the Quran portion
            file_name_list = str(theme).split(' ')
            file_name = '_'.join(file_name_list)
            content = f"""You are an expert Islamic story writer. Your job is to support the user in accomplishing the goal of writing stories for kids of 5 to 7 years of age on different verses or themes of the Quran. The stories should be between 250 to 300 words.
Follow these steps:
1.Write a motivating, inspiring story that is central to the theme or the message present in the verses
2.Proof read the story before producing the output
3.Add 3 questions in Aiken format that include right answer: 1 on the theme or the verse, 1 on story and 1 connecting theme and the story.
4.Produce the output in the English
Rules for creating the story:
1.If you find a matching story from hadith or tafseer, dramatize that particular story for kids.
2.Story must have an enticing title and it should be captivating the target audience of 5 to 7 years with a surprising plot twist or a climactic turn 
3.Use only children as main characters in the story. Use animals as only supporting characters to the main characters, only if it makes the story more interesting and captivating.
4.If possible, story must be in the form of dialogue or trialogue between the characters
5.Stories should avoid Disrespectful Terms, Sexual Content, Graphic Violence, Blasphemy, Misuse of Sacred Terms, Cultural Stereotypes, Sensitive Themes like violence, extremism, or terrorism, and misrepresentations that inaccurately represent Islamic teachings or beliefs
6.Every story should have a moral that is central to the theme mentioned.
Now, write a story on the theme : {theme}"""
            if custom_prompt:
                content = custom_prompt + f". Finally , write a story on the theme : {theme}"
            # Timer start
            start_time = time.time()
            # Generate story
            try:
                response = openai.ChatCompletion.create(
                    model=model,
                    messages=[
                        {"role": "system", "content": "You are an expert in Islamic story writing."},
                        {"role": "user", "content": content}
                    ]
                )
                story = response['choices'][0]['message']['content']
                st.write(f'Successfully created story from theme: "{theme}"')
            except:
                story = "Error in connection"
                st.write(f'Failed to create story from theme: {theme}"')
            # Calculate the time taken
            end_time = time.time()
            time_taken = end_time - start_time
            st.write(f"Time taken to generate story: {time_taken:.2f} seconds")
            
            story_list.append(story)

            # Save the story to a PDF file
            docx_file = save_word_docx(file_name, str(story))
            docx_files.append(docx_file)

            st.markdown(f"**Story {index+1}** created and saved as {docx_file}")            

        # Combine all stories into a single DataFrame
        df['Story'] = story_list
        st.write(df)
        # st.write("Creating a ZIP file...")

        # Create a zip file containing PDFs
        zip_buffer = BytesIO()
        with zipfile.ZipFile(zip_buffer, "w") as zf:
            for docx_file in docx_files:
                zf.write(docx_file, os.path.basename(docx_file))

        zip_file_name = "Stories.zip"
        zip_buffer.seek(0)
        code_end = time.time()
        total_code_time = code_end - code_start
        st.write(f"Time taken for the entire code to run: {total_code_time:.2f} seconds")
        

        # Add a download button for the zip file
        st.download_button(
            label="Download Stories as Zip",
            data=zip_buffer,
            key="download_stories_zip",
            file_name=zip_file_name,
            mime="application/zip"
        )

# Backend Code 2 -- End

input_mode = st.sidebar.radio("", ["Enter Quran Portion", "Upload Excel/CSV File"])

if input_mode == "Enter Quran Portion":
    gpt_version, df, custom_prompt = quran_portion()
else:
    gpt_version, df, custom_prompt = upload_and_process_file()

get_button = st.sidebar.button("Get Stories")

if get_button:
    output_code(gpt_version, df)
